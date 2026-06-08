"""
DOCX 格式化器

将 Thesis 对象（Markdown 章节 + 参考文献）转换为格式化的 .docx 文件。

功能：
1. Markdown → DOCX 转换（标题/正文/列表/表格）
2. 应用模板样式（页边距/字体/大小/对齐）
3. 插入参考文献列表
4. 注入目录域（TOC）
5. 输出完整论文 .docx 和大纲 .md

依赖：python-docx（基础操作）、markdown（Markdown 解析）
"""

import re
import os
import logging
from typing import List, Optional, Dict, Any
from datetime import datetime

from core.models import Thesis, Chapter, Reference, TemplateStyles, StyleDef, Outline

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class DocxFormatter:
    """DOCX 格式化器

    将 Thesis 对象的内容渲染为 .docx 文件。

    Attributes:
        thesis: 论文状态对象
    """

    def __init__(self):
        """初始化 DOCX 格式化器"""
        if not HAS_DOCX:
            logger.warning("python-docx 未安装，.docx 输出功能不可用")

    def create_document(self, thesis: Thesis, output_path: str) -> str:
        """创建完整的论文 .docx 文件

        流程：
        1. 创建空白文档 → 应用模板样式
        2. 插入封面页
        3. 插入目录域
        4. 逐章写入正文
        5. 插入参考文献
        6. 保存文件

        Args:
            thesis: 论文状态对象
            output_path: 输出文件路径

        Returns:
            str: 输出文件路径

        Raises:
            ImportError: python-docx 未安装
            RuntimeError: 文档创建失败
        """
        if not HAS_DOCX:
            raise ImportError("python-docx 未安装，无法生成 .docx 文件")

        try:
            doc = Document()

            # 1. 应用模板样式
            styles = thesis.template_styles or TemplateStyles()
            self._apply_page_setup(doc, styles)
            self._apply_styles(doc, styles)

            # 2. 封面页
            self._add_cover_page(doc, thesis)

            # 3. 分页 + 目录
            doc.add_page_break()
            self._add_toc(doc)

            # 4. 逐章写入正文
            for i, chapter in enumerate(thesis.chapters):
                if chapter.status in ("done", "edited") and chapter.content_markdown:
                    doc.add_page_break()
                    self._write_chapter(doc, chapter, i + 1, styles)

            # 4.5 P2: 图表附录（如果用户上传了数据）
            if thesis.user_data.get("raw_data"):
                try:
                    self._add_chart_appendix(doc, thesis)
                except Exception as e:
                    logger.warning(f"图表生成失败（非致命）: {e}")

            # 5. 参考文献
            if thesis.references:
                doc.add_page_break()
                self._write_references(doc, thesis.references, styles)

            # 6. 保存
            doc.save(output_path)
            logger.info(f"论文文档已保存: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"创建文档失败: {e}")
            raise RuntimeError(f"文档创建失败: {e}") from e

    def create_outline_doc(self, thesis: Thesis, output_path: str) -> str:
        """生成大纲 Markdown 文件

        Args:
            thesis: 论文状态对象
            output_path: 输出文件路径

        Returns:
            str: 输出文件路径
        """
        if thesis.outline:
            outline_md = thesis.outline.to_markdown()
        else:
            outline_md = "# 论文大纲\n\n（暂无大纲）"

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(outline_md)

        logger.info(f"大纲文件已保存: {output_path}")
        return output_path

    # ── 页面设置 ──────────────────────────────────────────────

    def _apply_page_setup(self, doc: "Document", styles: TemplateStyles) -> None:
        """应用页面边距设置

        Args:
            doc: Document 对象
            styles: 模板样式
        """
        for section in doc.sections:
            margins = styles.page_margins
            section.top_margin = Cm(margins.get("top", 25.4) / 10)
            section.bottom_margin = Cm(margins.get("bottom", 25.4) / 10)
            section.left_margin = Cm(margins.get("left", 31.7) / 10)
            section.right_margin = Cm(margins.get("right", 25.4) / 10)

    def _apply_styles(self, doc: "Document", styles: TemplateStyles) -> None:
        """应用正文默认样式

        Args:
            doc: Document 对象
            styles: 模板样式
        """
        if styles.body_style:
            bs = styles.body_style
            try:
                normal_style = doc.styles["Normal"]
                normal_style.font.name = bs.font_name
                normal_style.font.size = Pt(bs.font_size)
                normal_style.font.bold = bs.bold
                # 设置中文字体
                normal_style.element.rPr.rFonts.set(qn('w:eastAsia'), bs.font_name)
            except Exception as e:
                logger.debug(f"应用 Normal 样式失败: {e}")

    # ── 封面页 ────────────────────────────────────────────────

    def _add_cover_page(self, doc: "Document", thesis: Thesis) -> None:
        """添加论文封面页

        包含：论文标题、关键词、学科信息、生成信息、伦理声明。

        Args:
            doc: Document 对象
            thesis: 论文状态
        """
        # 空行
        for _ in range(5):
            doc.add_paragraph()

        # 论文标题
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(thesis.topic)
        title_run.bold = True
        title_run.font.size = Pt(22)
        title_run.font.name = "黑体"
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        doc.add_paragraph()

        # 副标题/信息
        info_items = [
            f"学科方向：{thesis.outline.discipline if thesis.outline else '——'}",
            f"关键词：{'、'.join(thesis.keywords)}",
            f"目标字数：{thesis.target_word_count:,} 字",
        ]
        for item in info_items:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(item)
            run.font.size = Pt(14)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 空行
        for _ in range(3):
            doc.add_paragraph()

        # 生成信息
        gen_para = doc.add_paragraph()
        gen_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        gen_run = gen_para.add_run(
            f"本文由 AI 辅助生成系统于 {datetime.now().strftime('%Y年%m月%d日')} 生成\n"
            "仅供学习研究参考 · 严禁直接作为学位论文提交"
        )
        gen_run.font.size = Pt(10)
        gen_run.font.color.rgb = RGBColor(128, 128, 128)

    # ── 目录 ──────────────────────────────────────────────────

    def _add_toc(self, doc: "Document") -> None:
        """插入目录域（TOC Field）

        生成的是域代码，在 Word 中打开后右键更新即可生成目录。

        Args:
            doc: Document 对象
        """
        # 目录标题
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("目  录")
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        doc.add_paragraph()

        # 插入 TOC 域
        para = doc.add_paragraph()
        run = para.add_run()
        # 构建 TOC 域 XML
        fldChar_begin = parse_xml(
            f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'
        )
        run._element.append(fldChar_begin)

        instrText = parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>'
        )
        run._element.append(instrText)

        fldChar_separate = parse_xml(
            f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'
        )
        run._element.append(fldChar_separate)

        # 提示文字
        hint_run = para.add_run("（请在 Word 中右键此处 → 更新域，以生成目录）")
        hint_run.font.size = Pt(10)
        hint_run.font.color.rgb = RGBColor(128, 128, 128)

        fldChar_end = parse_xml(
            f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'
        )
        run._element.append(fldChar_end)

    # ── 章节写入 ──────────────────────────────────────────────

    def _write_chapter(
        self,
        doc: "Document",
        chapter: Chapter,
        chapter_num: int,
        styles: TemplateStyles,
    ) -> None:
        """将单章内容写入文档

        Markdown 转换规则：
        - ## → 一级标题（章标题）
        - ### → 二级标题
        - #### → 三级标题
        - 普通段落 → 正文
        - - 开头 → 无序列表
        - 1. 开头 → 有序列表
        - 表格 → Word 表格

        Args:
            doc: Document 对象
            chapter: 章节对象
            chapter_num: 章节序号
            styles: 模板样式
        """
        content = chapter.content_markdown
        if not content:
            return

        lines = content.split("\n")

        i = 0
        while i < len(lines):
            line = lines[i]

            # 跳过空行
            if not line.strip():
                i += 1
                continue

            stripped = line.strip()

            # 水平线
            if re.match(r'^[-*_]{3,}$', stripped):
                doc.add_paragraph()
                i += 1
                continue

            # P2: 公式（$$...$$）
            formula_match = re.match(r'^\$\$(.+?)\$\$', stripped)
            if formula_match:
                self._add_formula_image(doc, formula_match.group(1).strip())
                i += 1
                continue

            # 标题
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
            if heading_match:
                level_from_md = len(heading_match.group(1))
                title_text = self._clean_markdown_inline(heading_match.group(2))
                # 映射：## → level1, ### → level2, #### → level3
                docx_level = level_from_md - 1
                self._add_heading(doc, title_text, min(docx_level, 3), styles)
                i += 1
                continue

            # 无序列表
            if re.match(r'^[\s]*[-*+]\s+', stripped):
                list_items = []
                while i < len(lines) and re.match(r'^[\s]*[-*+]\s+', lines[i].strip()):
                    item_text = re.sub(r'^[\s]*[-*+]\s+', '', lines[i].strip())
                    list_items.append(self._clean_markdown_inline(item_text))
                    i += 1
                for item in list_items:
                    self._add_list_item(doc, item, bullet=True)
                continue

            # 有序列表
            if re.match(r'^[\s]*\d+[.)]\s+', stripped):
                list_items = []
                while i < len(lines) and re.match(r'^[\s]*\d+[.)]\s+', lines[i].strip()):
                    item_text = re.sub(r'^[\s]*\d+[.)]\s+', '', lines[i].strip())
                    list_items.append(self._clean_markdown_inline(item_text))
                    i += 1
                for idx, item in enumerate(list_items, 1):
                    self._add_list_item(doc, item, bullet=False, number=idx)
                continue

            # 表格（检测 | 分隔的行）
            if '|' in stripped and i + 1 < len(lines) and '|---' in lines[i + 1]:
                table_data = self._parse_markdown_table(lines, i)
                if table_data:
                    self._add_table(doc, table_data[0], table_data[1])
                    i = table_data[2]  # 跳到表格之后
                    continue

            # 引用块
            if stripped.startswith('>'):
                quote_text = stripped[1:].strip()
                quote_lines = [quote_text]
                i += 1
                while i < len(lines) and lines[i].strip().startswith('>'):
                    quote_lines.append(lines[i].strip()[1:].strip())
                    i += 1
                for ql in quote_lines:
                    self._add_quote(doc, ql)
                continue

            # 普通段落
            para_text = self._clean_markdown_inline(stripped)
            self._add_paragraph(doc, para_text, styles)
            i += 1

    # ── 参考文献 ──────────────────────────────────────────────

    def _add_chart_appendix(self, doc: "Document", thesis: "Thesis") -> None:
        """P2: 生成图表附录

        Args:
            doc: Document 对象
            thesis: 论文状态
        """
        from core.charts import ChartGenerator, HAS_MPL
        if not HAS_MPL:
            logger.warning("matplotlib 不可用，跳过图表生成")
            return

        try:
            cg = ChartGenerator()
            data_dict = {
                "raw_data": thesis.user_data.get("raw_data", []),
                "columns": thesis.user_data.get("data_columns", []),
            }
            result = cg.try_generate_from_data(data_dict)
            if result is None:
                return

            buf, title = result

            # 图表页
            doc.add_page_break()
            heading = doc.add_heading("附录：数据分析图表", level=2)
            self._apply_heading_style(heading, styles=None)

            para = doc.add_paragraph(f"\n{title}\n")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 保存临时图片并插入
            import tempfile
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                tmp.write(buf.getvalue())
                tmp.flush()
                doc.add_picture(tmp.name, width=Inches(5.5))
                os.unlink(tmp.name)

            logger.info("图表附录已添加")

        except Exception as e:
            logger.warning(f"图表生成失败: {e}")

    def _write_references(
        self,
        doc: "Document",
        references: List[Reference],
        styles: TemplateStyles,
    ) -> None:
        """写入参考文献列表

        按 GB/T 7714 格式排列。

        Args:
            doc: Document 对象
            references: 参考文献列表
            styles: 模板样式
        """
        # 标题
        self._add_heading(doc, "参考文献", 1, styles)

        for i, ref in enumerate(references, 1):
            citation = ref.to_gb7714()
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.left_indent = Cm(0.74)

            # 编号
            num_run = para.add_run(f"[{i}] ")
            num_run.font.size = Pt(10.5)

            # 内容
            text_run = para.add_run(citation)
            text_run.font.size = Pt(10.5)
            text_run.font.name = "宋体"
            text_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ── 元素级写入方法 ────────────────────────────────────────

    def _add_heading(
        self,
        doc: "Document",
        text: str,
        level: int,
        styles: TemplateStyles,
    ) -> None:
        """添加标题段落

        Args:
            doc: Document 对象
            text: 标题文本
            level: 标题级别（1/2/3）
            styles: 模板样式
        """
        para = doc.add_paragraph()

        if level in styles.heading_styles:
            hs = styles.heading_styles[level]
            if hs.alignment == 1:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif hs.alignment == 2:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            run = para.add_run(text)
            run.bold = hs.bold
            run.font.size = Pt(hs.font_size)
            run.font.name = hs.font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), hs.font_name)
        else:
            # 默认
            sizes = {1: 16, 2: 14, 3: 12}
            run = para.add_run(text)
            run.bold = True
            run.font.size = Pt(sizes.get(level, 12))
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # 设置段前段后间距
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)

    def _add_paragraph(
        self,
        doc: "Document",
        text: str,
        styles: TemplateStyles,
    ) -> None:
        """添加正文段落

        Args:
            doc: Document 对象
            text: 段落文本
            styles: 模板样式
        """
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进两字符
        para.paragraph_format.line_spacing = 1.5

        run = para.add_run(text)
        if styles.body_style:
            bs = styles.body_style
            run.font.size = Pt(bs.font_size)
            run.font.name = bs.font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), bs.font_name)
        else:
            run.font.size = Pt(12)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def _add_list_item(
        self,
        doc: "Document",
        text: str,
        bullet: bool = True,
        number: Optional[int] = None,
    ) -> None:
        """添加列表项

        Args:
            doc: Document 对象
            text: 列表项文本
            bullet: True=无序, False=有序
            number: 有序列表编号
        """
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.74)
        para.paragraph_format.first_line_indent = Cm(0)

        if bullet:
            prefix = "• "
        else:
            prefix = f"{number}. " if number else ""

        run = para.add_run(f"{prefix}{text}")
        run.font.size = Pt(12)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def _add_quote(self, doc: "Document", text: str) -> None:
        """添加引用块

        Args:
            doc: Document 对象
            text: 引用文本
        """
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(1.5)
        para.paragraph_format.right_indent = Cm(1.0)

        run = para.add_run(text)
        run.font.size = Pt(10.5)
        run.font.italic = True
        run.font.color.rgb = RGBColor(80, 80, 80)

    def _add_table(self, doc: "Document", headers: List[str], rows: List[List[str]]) -> None:
        """添加三线表（P1 升级：顶线粗/栏目线细/底线粗）

        符合学术规范：无左右竖线、无表内横线。

        Args:
            doc: Document 对象
            headers: 表头
            rows: 数据行
        """
        if not headers:
            return

        table = doc.add_table(rows=1 + len(rows), cols=len(headers))

        # ── 禁用所有默认边框 ──
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                # 清除单元格边框
                tcBorders = parse_xml(
                    f'<w:tcBorders {nsdecls("w")}>'
                    f'<w:top w:val="nil"/>'
                    f'<w:left w:val="nil"/>'
                    f'<w:bottom w:val="nil"/>'
                    f'<w:right w:val="nil"/>'
                    f'</w:tcBorders>'
                )
                tcPr.append(tcBorders)

        # ── 表头行 ──
        for j, header in enumerate(headers):
            cell = table.rows[0].cells[j]
            self._set_cell_text(cell, header, bold=True, alignment="center", font_size=10)

        # ── 数据行 ──
        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                if j < len(headers):
                    cell = table.rows[i + 1].cells[j]
                    self._set_cell_text(cell, str(cell_text), bold=False, alignment="center", font_size=10)

        # ── 绘制三线 ──
        # 顶线（粗 1.5pt）
        self._set_row_border(table.rows[0], "top", sz="12")
        # 栏目线（细 0.75pt，表头下方）
        self._set_row_border(table.rows[0], "bottom", sz="6")
        # 底线（粗 1.5pt，表格最后一行）
        if len(rows) > 0:
            self._set_row_border(table.rows[-1], "bottom", sz="12")

        doc.add_paragraph()  # 表格后空行

    @staticmethod
    def _set_cell_text(cell, text: str, bold: bool, alignment: str = "center", font_size: int = 10) -> None:
        """设置单元格文本和格式

        Args:
            cell: Word 单元格对象
            text: 文本内容
            bold: 是否加粗
            alignment: 对齐方式
            font_size: 字号
        """
        for para in cell.paragraphs:
            para.alignment = {
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
            }.get(alignment, WD_ALIGN_PARAGRAPH.CENTER)
            # 清除默认空段落内容
            if para.runs:
                para.clear()
            run = para.add_run(text)
            run.bold = bold
            run.font.size = Pt(font_size)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    @staticmethod
    def _set_row_border(row, edge: str, sz: str) -> None:
        """为行设置指定边框

        Args:
            row: Word 行对象
            edge: 边框位置 ("top" / "bottom")
            sz: 线宽 ("12"=1.5pt, "6"=0.75pt)
        """
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
                tcPr.append(tcBorders)
            edge_elem = tcBorders.find(qn(f'w:{edge}'))
            if edge_elem is not None:
                tcBorders.remove(edge_elem)
            new_edge = parse_xml(
                f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="{sz}" '
                f'w:space="0" w:color="000000"/>'
            )
            tcBorders.append(new_edge)

    # ── Markdown 工具 ─────────────────────────────────────────

    def _add_formula_image(self, doc: "Document", formula: str) -> None:
        """P2: 渲染并插入公式图片

        Args:
            doc: Document 对象
            formula: LaTeX 公式文本
        """
        try:
            from core.formula_renderer import FormulaRenderer, HAS_MPL
            if not HAS_MPL:
                # 回退：纯文本显示
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(f"$$ {formula} $$")
                run.italic = True
                run.font.size = Pt(11)
                return

            fr = FormulaRenderer()
            buf = fr.render_display(formula)

            import tempfile
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                tmp.write(buf.getvalue())
                tmp.flush()
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(tmp.name, width=Inches(4.5))
                os.unlink(tmp.name)

        except Exception as e:
            logger.debug(f"公式渲染失败，回退纯文本: {e}")
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run(f"[公式: {formula[:50]}]")

    @staticmethod
    def _clean_markdown_inline(text: str) -> str:
        """清理 Markdown 行内格式标记

        Args:
            text: Markdown 文本

        Returns:
            str: 纯文本
        """
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'`(.+?)`', r'\1', text)
        text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
        text = re.sub(r'!\[.*?\]\(.+?\)', '', text)
        return text.strip()

    @staticmethod
    def _parse_markdown_table(lines: List[str], start_index: int) -> Optional[tuple]:
        """解析 Markdown 表格

        格式：
        | A | B | C |
        |---|---|
        | 1 | 2 | 3 |

        Args:
            lines: 所有行
            start_index: 表格起始行索引

        Returns:
            Optional[tuple]: (headers, rows, next_index) 或 None
        """
        try:
            # 表头行
            header_line = lines[start_index].strip()
            headers = [h.strip() for h in header_line.split('|') if h.strip()]

            # 分隔行（跳过）
            if start_index + 1 >= len(lines):
                return None

            # 数据行
            rows = []
            idx = start_index + 2
            while idx < len(lines):
                line = lines[idx].strip()
                if '|' not in line:
                    break
                cells = [c.strip() for c in line.split('|') if c.strip()]
                if cells:
                    rows.append(cells)
                idx += 1

            return (headers, rows, idx)
        except Exception:
            return None
