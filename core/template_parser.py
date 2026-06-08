"""
模板解析器

解析用户上传的 .docx 模板文件，提取：
- 页面设置（页边距、纸张大小）
- 标题样式（字体、大小、加粗、对齐方式）
- 正文样式
- 页眉页脚信息

未上传模板时使用默认的学术论文格式。
"""

import os
import logging
from typing import Optional, Dict, Any

from core.models import TemplateStyles, StyleDef

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Inches, Cm, Pt, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class TemplateParser:
    """模板解析器

    从 .docx 模板中提取样式信息，生成 TemplateStyles 对象。
    解析失败时回退到默认样式。

    Attributes:
        default_styles: 默认的 TemplateStyles（学术论文标准格式）
    """

    def __init__(self):
        """初始化模板解析器"""
        if not HAS_DOCX:
            logger.warning("python-docx 未安装，模板解析将使用默认样式")

        # 预置默认样式
        self.default_styles = TemplateStyles()
        self._presets = self._build_presets()

    # ── P2 新增：预设模板注册表 ──────────────────────────────

    @staticmethod
    def _build_presets() -> Dict[str, TemplateStyles]:
        """构建预设模板注册表"""
        presets = {}
        ql = TemplateStyles(
            page_margins={"top": 25.4, "bottom": 25.4, "left": 31.7, "right": 25.4},
            heading_styles={
                1: StyleDef(name="标题1", font_name="黑体", font_size=16, bold=True, alignment=1),
                2: StyleDef(name="标题2", font_name="黑体", font_size=14, bold=True, alignment=0),
                3: StyleDef(name="标题3", font_name="黑体", font_size=12, bold=True, alignment=0),
            },
            body_style=StyleDef(name="正文", font_name="宋体", font_size=12, bold=False, alignment=0),
            template_name="齐鲁理工学院标准",
            template_description="毕业论文标准格式：章标题黑体16pt居中，节标题黑体14pt左对齐，正文宋体12pt",
        )
        presets["qlu"] = ql
        academic = TemplateStyles(
            page_margins={"top": 25.4, "bottom": 25.4, "left": 25.4, "right": 25.4},
            heading_styles={
                1: StyleDef(name="标题1", font_name="Times New Roman", font_size=14, bold=True, alignment=0),
                2: StyleDef(name="标题2", font_name="Times New Roman", font_size=12, bold=True, alignment=0),
                3: StyleDef(name="标题3", font_name="Times New Roman", font_size=12, bold=True, alignment=0),
            },
            body_style=StyleDef(name="正文", font_name="Times New Roman", font_size=12, bold=False, alignment=3),
            template_name="通用英文学术",
            template_description="英文论文格式：Times New Roman 12pt，标题加粗左对齐，正文两端对齐",
        )
        presets["academic"] = academic
        engr = TemplateStyles(
            page_margins={"top": 20.0, "bottom": 20.0, "left": 30.0, "right": 25.0},
            heading_styles={
                1: StyleDef(name="标题1", font_name="黑体", font_size=15, bold=True, alignment=1),
                2: StyleDef(name="标题2", font_name="黑体", font_size=13, bold=True, alignment=0),
                3: StyleDef(name="标题3", font_name="宋体", font_size=12, bold=True, alignment=0),
            },
            body_style=StyleDef(name="正文", font_name="宋体", font_size=11, bold=False, alignment=0),
            template_name="工科规范",
            template_description="工程类论文：页边距缩小，正文宋体11pt，适合图表较多的论文",
        )
        presets["engineering"] = engr
        return presets

    def get_preset_list(self) -> list:
        """获取预设模板列表"""
        return [
            {"id": tid, "name": ts.template_name, "description": ts.template_description}
            for tid, ts in self._presets.items()
        ]

    def get_preset(self, template_id: str) -> Optional[TemplateStyles]:
        """通过 ID 获取预设模板"""
        return self._presets.get(template_id)

    # ── 原有解析方法 ──────────────────────────────────────────

    def parse(self, docx_path: str) -> TemplateStyles:
        """解析 .docx 模板文件

        提取页面设置和样式信息。如果文件不存在或解析失败，
        返回默认样式。

        Args:
            docx_path: .docx 文件路径（用户上传的模板）

        Returns:
            TemplateStyles: 提取的样式信息或默认样式
        """
        if not HAS_DOCX:
            logger.warning("python-docx 未安装，返回默认样式")
            return self.default_styles

        # 文件不存在时返回默认样式
        if not docx_path or not os.path.isfile(docx_path):
            logger.info(f"模板文件不存在 '{docx_path}'，使用默认样式")
            return self.default_styles

        try:
            doc = Document(docx_path)
            styles = TemplateStyles()

            # 1. 提取页边距
            styles.page_margins = self._extract_page_margins(doc)

            # 2. 提取标题样式
            styles.heading_styles = self._extract_heading_styles(doc)

            # 3. 提取正文样式
            styles.body_style = self._extract_body_style(doc)

            # 4. 提取页眉页脚
            styles.header_footer = self._extract_header_footer(doc)

            logger.info(
                f"模板解析成功: '{os.path.basename(docx_path)}', "
                f"提取了 {len(styles.heading_styles)} 种标题样式"
            )
            return styles

        except Exception as e:
            logger.error(f"模板解析失败: {e}，回退到默认样式")
            return self.default_styles

    # ── 内部提取方法 ──────────────────────────────────────────

    def _extract_page_margins(self, doc: "Document") -> Dict[str, float]:
        """提取页面边距

        Args:
            doc: python-docx Document 对象

        Returns:
            Dict: 页边距字典（单位 mm）
        """
        try:
            section = doc.sections[0]
            # python-docx 内部使用 Emu，转换为 mm
            margins = {
                "top": round(section.top_margin / 360000, 1) if section.top_margin else 25.4,
                "bottom": round(section.bottom_margin / 360000, 1) if section.bottom_margin else 25.4,
                "left": round(section.left_margin / 360000, 1) if section.left_margin else 31.7,
                "right": round(section.right_margin / 360000, 1) if section.right_margin else 25.4,
            }
            return margins
        except Exception as e:
            logger.debug(f"提取页边距失败: {e}")
            return {"top": 25.4, "bottom": 25.4, "left": 31.7, "right": 25.4}

    def _extract_heading_styles(self, doc: "Document") -> Dict[int, StyleDef]:
        """提取标题样式

        遍历所有段落，找到使用 Heading 样式的段落并提取格式。

        Args:
            doc: python-docx Document 对象

        Returns:
            Dict[int, StyleDef]: level → StyleDef 映射
        """
        styles: Dict[int, StyleDef] = {}
        heading_names = {
            1: ["Heading 1", "heading 1", "标题 1", "标题1", "1. heading 1"],
            2: ["Heading 2", "heading 2", "标题 2", "标题2", "2. heading 2"],
            3: ["Heading 3", "heading 3", "标题 3", "标题3", "3. heading 3"],
        }

        for level, names in heading_names.items():
            try:
                style = None
                for name in names:
                    try:
                        style = doc.styles[name]
                        break
                    except KeyError:
                        continue

                if style and style.paragraph_format:
                    font = style.font
                    style_def = StyleDef(
                        name=f"标题{level}",
                        font_name=self._get_font_name(font),
                        font_size=self._get_font_size_pt(font),
                        bold=font.bold or False,
                        alignment=self._get_alignment(style.paragraph_format.alignment),
                    )
                    styles[level] = style_def
                else:
                    # 使用默认
                    styles[level] = self.default_styles.heading_styles.get(level, StyleDef())

            except Exception as e:
                logger.debug(f"提取标题 level={level} 样式失败: {e}")
                styles[level] = self.default_styles.heading_styles.get(level, StyleDef())

        return styles

    def _extract_body_style(self, doc: "Document") -> Optional[StyleDef]:
        """提取正文样式

        从 Normal 样式中提取。

        Args:
            doc: python-docx Document 对象

        Returns:
            Optional[StyleDef]: 正文样式定义
        """
        try:
            style = doc.styles["Normal"]
            font = style.font
            return StyleDef(
                name="正文",
                font_name=self._get_font_name(font),
                font_size=self._get_font_size_pt(font),
                bold=False,
                alignment=0,  # 正文默认左对齐
            )
        except Exception as e:
            logger.debug(f"提取正文样式失败: {e}")
            return StyleDef()

    def _extract_header_footer(self, doc: "Document") -> Optional[Dict[str, Any]]:
        """提取页眉页脚信息

        Args:
            doc: python-docx Document 对象

        Returns:
            Optional[Dict]: 页眉页脚文本信息
        """
        try:
            section = doc.sections[0]
            header_text = ""
            footer_text = ""

            # 提取页眉文本
            header = section.header
            if header and not header.is_linked_to_previous:
                for para in header.paragraphs:
                    header_text += para.text + "\n"

            # 提取页脚文本
            footer = section.footer
            if footer and not footer.is_linked_to_previous:
                for para in footer.paragraphs:
                    footer_text += para.text + "\n"

            if header_text.strip() or footer_text.strip():
                return {
                    "header": header_text.strip(),
                    "footer": footer_text.strip(),
                }
            return None
        except Exception as e:
            logger.debug(f"提取页眉页脚失败: {e}")
            return None

    # ── 辅助方法 ──────────────────────────────────────────────

    @staticmethod
    def _get_font_name(font) -> str:
        """安全获取字体名称

        Args:
            font: python-docx Font 对象

        Returns:
            str: 字体名
        """
        try:
            return font.name or "宋体"
        except Exception:
            return "宋体"

    @staticmethod
    def _get_font_size_pt(font) -> int:
        """安全获取字体大小（pt）

        Args:
            font: python-docx Font 对象

        Returns:
            int: 字体大小（pt）
        """
        try:
            if font.size:
                return int(font.size / 12700)  # Emu → pt
            return 12
        except Exception:
            return 12

    @staticmethod
    def _get_alignment(alignment) -> int:
        """转换对齐方式为整数编码

        Args:
            alignment: WD_ALIGN_PARAGRAPH 枚举值

        Returns:
            int: 0=左, 1=中, 2=右
        """
        if alignment is None:
            return 0
        try:
            mapping = {
                1: 1,  # CENTER
                2: 2,  # RIGHT
                3: 0,  # JUSTIFY → 左对齐
                0: 0,  # LEFT
            }
            return mapping.get(int(alignment), 0)
        except Exception:
            return 0
