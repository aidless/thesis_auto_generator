"""v2.1 第三次实测 — A+B 效果验证"""
import os, time, re, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from config import load_config
from llm.base import create_llm_client
from core.outline_generator import OutlineGenerator
from core.chapter_generator import ChapterGenerator
from core.reference_fetcher import ReferenceFetcher
from core.docx_formatter import DocxFormatter
from core.models import Thesis, TemplateStyles
from utils.watermark import add_watermark_and_disclaimer
from docx import Document

config = load_config()
llm = create_llm_client("deepseek", config)
os.makedirs("data/output", exist_ok=True)

topic = "基于Spring Boot的码头船只出行管理系统设计与实现"
keywords = ["Spring Boot", "码头管理", "船只调度", "Java"]

t0 = time.time()

# Phase 1
print("Phase 1: 大纲...")
og = OutlineGenerator(llm)
outline = og.generate(topic, keywords, "软件工程")
chapters_nodes = outline.get_chapters()
print(f"  {len(chapters_nodes)} 章")

# Phase 2
print("Phase 2: 逐章生成 (v2.1 A+B)...")
cg = ChapterGenerator(llm)
all_chapters = []
chapter_stats = []
for i, node in enumerate(chapters_nodes):
    ch_start = time.time()
    ch = cg.generate_chapter(outline, node, all_chapters)
    all_chapters.append(ch)
    elapsed = time.time() - ch_start

    content = ch.content_markdown
    has_table = "|" in content and "---" in content
    has_reasoning = bool(re.search(r"(因为|原因在于|选择.*理由|之所以|是由于|综合考虑)", content))
    has_empty = bool(re.search(r"(具有重要意义|取得了良好效果|得到了广泛应用|十分必要)", content))

    flags = ""
    if has_table:
        flags += " TABLE"
    if has_reasoning:
        flags += " REASONING"
    if has_empty:
        flags += " EMPTY-PHRASE"

    chapter_stats.append({
        "chapter": node.title[:20], "words": ch.word_count, "time": elapsed,
        "has_table": has_table, "has_reasoning": has_reasoning, "has_empty": has_empty,
    })
    print(f"  [{i+1}/{len(chapters_nodes)}] {node.title[:20]:20s} {ch.word_count:5d}字 {elapsed:3.0f}s{flags}")

total_words = sum(ch.word_count for ch in all_chapters)

# Phase 3
print("Phase 3: 参考文献...")
rf = ReferenceFetcher()
refs = rf.fetch_by_keywords(keywords, limit=15, topic=topic)
sb_refs = [r for r in refs if "spring" in (r.title + r.abstract).lower()
           or "boot" in (r.title + r.abstract).lower()
           or "web" in r.title.lower()[:20]]
print(f"  {len(refs)} 篇 (其中 {len(sb_refs)} 篇与 Spring Boot/Web 直接相关)")

# Phase 4
print("Phase 4: DOCX...")
thesis = Thesis(topic=topic, keywords=keywords,
    target_word_count=config.target_word_count, outline=outline,
    chapters=all_chapters, references=refs, template_styles=TemplateStyles())
output = "data/output/real_e2e_v3.docx"
DocxFormatter().create_document(thesis, output)
final = "data/output/real_e2e_v3_final.docx"
add_watermark_and_disclaimer(output, final)

doc = Document(final)
tables = doc.tables
full_text = "\n".join(p.text for p in doc.paragraphs)

reasoning_count = len(re.findall(r"(因为|原因在于|选择.*理由|之所以|是由于|综合考虑)", full_text))
empty_count = len(re.findall(r"(具有重要意义|取得了良好效果|得到了广泛应用|十分必要)", full_text))
table_mentions = len(re.findall(r"(如表|见表|表\d|Table)", full_text))

total_time = time.time() - t0

# Print comparison
sep = "=" * 70
print()
# Compute metrics
ct = sum(1 for s in chapter_stats if s["has_table"])
cr = sum(1 for s in chapter_stats if s["has_reasoning"])
ce = sum(1 for s in chapter_stats if s["has_empty"])
n = len(chapter_stats)

print(sep)
print("  Third E2E: v2.1 (A+B Improvements)")
print(sep)
print(f"  {'Metric':<30} {'v2.0 (2nd run)':<20} {'v2.1 (this run)':<20}")
print(f"  {'-'*30} {'-'*20} {'-'*20}")
print(f"  {'Total Words':<30} {'17,486':<20} {total_words:<20,}")
print(f"  {'Total Time':<30} {'207s':<20} {total_time:.0f}s")
print(f"  {'Chapters with Tables':<30} {'4/7':<20} {ct}/{n}")
print(f"  {'Chapters with Reasoning':<30} {'N/A':<20} {cr}/{n}")
print(f"  {'Chapters with Empty Phrases':<30} {'N/A':<20} {ce}/{n}")
print(f"  {'DOCX Tables':<30} {'1':<20} {len(tables)}")
print(f"  {'Spring Boot refs':<30} {'N/A':<20} {len(sb_refs)}")
print(f"  {'Total reasoning phrases':<30} {'N/A':<20} {reasoning_count}")
print(f"  {'Total empty phrases':<30} {'N/A':<20} {empty_count}")
print(f"  {'Table references (as Table X)':<30} {'N/A':<20} {table_mentions}")
print(f"  {'File':<30} {'real_e2e_v2_final.docx':<20} real_e2e_v3_final.docx")
print(sep)
