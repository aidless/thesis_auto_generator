"""
Mock 端到端测试（test_e2e_mock.py）

用预制 JSON 模拟 LLM 响应，验证核心生成链路：
主题 → 大纲 → 逐章生成 → DOCX 输出

不需要 API Key，10 秒内完成全链路验证。
"""

import os
import sys
import json
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch, MagicMock
from typing import List

# 确保项目根目录在 sys.path
PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from core.models import (
    Thesis, Outline, OutlineNode, Chapter, Reference,
    TemplateStyles, StyleDef,
)
from core.template_parser import TemplateParser
from core.docx_formatter import DocxFormatter
from utils.file_utils import generate_output_filename, get_output_dir, ensure_dir
from utils.text_utils import estimate_word_count
from utils.watermark import add_watermark_and_disclaimer
from docx import Document  # P4: 用于 DOCX 结构验证


# ── 预制 LLM 响应（Mock 数据）─────────────────────────────

MOCK_OUTLINE_MD = """## 第一章 绪论
> 介绍研究背景、现状、目的和意义

### 1.1 研究背景与意义
### 1.2 国内外研究现状
### 1.3 主要研究内容
### 1.4 论文组织结构

## 第二章 相关理论与技术基础
> 介绍所用到的技术基础

### 2.1 Spring Boot 框架概述
### 2.2 Vue.js 前端框架
### 2.3 MySQL 数据库
### 2.4 本章小结

## 第三章 系统需求分析
> 分析系统需求

### 3.1 业务需求分析
### 3.2 功能需求分析
### 3.3 非功能需求分析
### 3.4 可行性分析

## 第四章 系统设计
> 详细设计系统架构和数据库

### 4.1 系统总体架构设计
### 4.2 功能模块设计
### 4.3 数据库设计
### 4.4 接口设计

## 第五章 系统实现与测试
> 展示系统实现和测试结果

### 5.1 开发环境配置
### 5.2 核心功能实现
### 5.3 系统测试
### 5.4 测试结果分析

## 第六章 总结与展望
> 总结工作并展望未来

### 6.1 工作总结
### 6.2 不足与展望
"""

MOCK_CHAPTER_CONTENTS = [
    # 第一章：绪论
    """## 第一章 绪论

### 1.1 研究背景与意义

随着信息技术的迅猛发展，码头船只出行管理面临着新的挑战和机遇。传统的管理方式依赖人工记录和纸质单据，效率低下且容易出错。开发一套基于 Spring Boot 的码头船只出行管理系统，能够实现船只调度的自动化、信息化和智能化，对于提升码头运营效率具有重要意义。

### 1.2 国内外研究现状

近年来，国内外学者在码头管理系统领域开展了大量研究。Smith 等人（2020）提出了一种基于物联网的智能码头调度方案，实现了船只到港的自动识别和泊位分配。李明等（2021）研究了基于微服务架构的港口管理系统，提高了系统的可扩展性和维护性。

| 研究者 | 年份 | 方法 | 局限 |
|--------|------|------|------|
| Smith et al. | 2020 | IoT + 自动调度 | 仅适用于小型码头 |
| 李明等 | 2021 | 微服务架构 | 部署复杂度高 |

### 1.3 主要研究内容

本文主要研究内容包括：码头船只出行管理系统的需求分析、系统架构设计、数据库设计、核心功能实现以及系统测试。

### 1.4 论文组织结构

本文共分为六章。第一章介绍研究背景和现状；第二章介绍相关技术；第三章进行需求分析；第四章阐述系统设计；第五章展示系统实现与测试；第六章进行总结与展望。""",

    # 第二章：技术基础
    """## 第二章 相关理论与技术基础

### 2.1 Spring Boot 框架概述

Spring Boot 是基于 Spring 框架的快速开发框架，它简化了 Spring 应用的初始搭建和开发过程。通过自动配置和起步依赖，开发者可以快速构建生产级别的独立应用。本文采用 Spring Boot 2.7.x 版本作为后端开发框架。

### 2.2 Vue.js 前端框架

Vue.js 是一套用于构建用户界面的渐进式 JavaScript 框架。它采用组件化的开发方式，配合 Element UI 组件库，可以快速构建美观的管理界面。本文前端采用 Vue 3 + Element Plus 技术栈。

### 2.3 MySQL 数据库

MySQL 是目前最流行的开源关系型数据库管理系统。它具有高性能、高可靠性和易用性等特点。本文使用 MySQL 8.0 作为数据持久化方案。

### 2.4 本章小结

本章介绍了码头船只出行管理系统所涉及的核心技术，包括 Spring Boot 后端框架、Vue.js 前端框架和 MySQL 数据库，为后续章节的系统设计与实现奠定了技术基础。""",

    # 第三章
    """## 第三章 系统需求分析

### 3.1 业务需求分析

码头船只出行管理系统面向码头管理人员和船只操作人员，核心业务包括船只登记、泊位分配、进出港记录和费用结算等。

### 3.2 功能需求分析

系统主要功能模块包括：用户管理模块、船只管理模块、泊位管理模块、调度管理模块和报表统计模块。

### 3.3 非功能需求分析

系统需满足响应时间小于 2 秒、并发用户数不低于 50 人、数据存储安全可靠等非功能性要求。

### 3.4 可行性分析

从技术可行性、经济可行性和操作可行性三个角度分析，本系统的开发是可行的。""",

    # 第四章
    """## 第四章 系统设计

### 4.1 系统总体架构设计

本系统采用 B/S 三层架构，前端使用 Vue.js 构建用户界面，后端使用 Spring Boot 提供 RESTful API，数据库使用 MySQL 存储业务数据。

### 4.2 功能模块设计

系统包含五大功能模块：用户管理、船只管理、泊位管理、调度管理和报表统计。

### 4.3 数据库设计

数据库包含用户表、船只表、泊位表、调度记录表和费用表等核心数据表。采用第三范式进行规范化设计，确保数据一致性。

### 4.4 接口设计

系统采用 RESTful 风格设计 API 接口，使用 JSON 格式进行数据交换，通过 JWT 令牌实现身份认证。""",

    # 第五章
    """## 第五章 系统实现与测试

### 5.1 开发环境配置

开发环境：Windows 11 操作系统，JDK 11，IntelliJ IDEA 2023，MySQL 8.0，Node.js 18。

### 5.2 核心功能实现

实现了用户登录注册、船只信息管理、泊位状态管理、调度任务分配等核心功能。

### 5.3 系统测试

采用黑盒测试和白盒测试相结合的方法，编写了 50 个测试用例，覆盖了主要功能模块。

| 测试类型 | 用例数 | 通过率 |
|----------|--------|--------|
| 单元测试 | 30 | 100% |
| 集成测试 | 15 | 93.3% |
| 系统测试 | 5 | 100% |

### 5.4 测试结果分析

测试结果表明，系统功能完整、性能稳定，满足设计需求。""",

    # 第六章
    """## 第六章 总结与展望

### 6.1 工作总结

本文设计并实现了一套基于 Spring Boot 的码头船只出行管理系统。系统采用 B/S 架构，Spring Boot 作为后端框架，Vue.js 作为前端框架，MySQL 作为数据库。

### 6.2 不足与展望

系统在智能调度算法方面还有优化空间，未来可以引入机器学习算法实现更高效的泊位分配。同时可以考虑增加移动端应用，方便现场操作人员使用。""",
]

MOCK_REFERENCES = [
    Reference(
        key="Smith2020",
        title="Intelligent Dock Scheduling Using IoT Technology",
        authors=["Smith J.", "Brown A."],
        year=2020,
        venue="Journal of Maritime Technology",
        doi="10.1000/jmt.2020.001",
        citation_text="Smith J., Brown A. Intelligent Dock Scheduling Using IoT Technology[J]. Journal of Maritime Technology, 2020.",
    ),
    Reference(
        key="Li2021",
        title="微服务架构在港口管理系统中的应用研究",
        authors=["李明", "王伟"],
        year=2021,
        venue="计算机应用研究",
        citation_text="李明, 王伟. 微服务架构在港口管理系统中的应用研究[J]. 计算机应用研究, 2021.",
    ),
    Reference(
        key="Zhang2022",
        title="基于 Spring Boot 的企业级应用快速开发方法",
        authors=["张华"],
        year=2022,
        venue="软件工程",
        citation_text="张华. 基于 Spring Boot 的企业级应用快速开发方法[J]. 软件工程, 2022.",
    ),
]


class MockLLMClient:
    """Mock LLM 客户端"""

    def __init__(self):
        self.max_tokens = 4096
        self._chapter_index = 0  # 追踪当前生成到第几章

    def chat(self, messages, **kwargs):
        """根据 prompt 内容返回对应 Mock 响应"""
        for msg in messages:
            content = msg.get("content", "")
            if "大纲" in content and "章节" in content:
                return MOCK_OUTLINE_MD
            if "第" in content and "章" in content and "撰写" in content:
                # 按顺序返回各章内容
                idx = self._chapter_index
                if idx < len(MOCK_CHAPTER_CONTENTS):
                    self._chapter_index += 1
                    return MOCK_CHAPTER_CONTENTS[idx]
        # 默认返回
        return MOCK_CHAPTER_CONTENTS[0]


class TestE2EMock(unittest.TestCase):
    """Mock 端到端测试套件"""

    @classmethod
    def setUpClass(cls):
        """初始化测试环境"""
        cls.temp_dir = tempfile.mkdtemp(prefix="thesis_test_")
        cls.output_dir = Path(cls.temp_dir) / "output"
        cls.output_dir.mkdir(exist_ok=True)

    def test_01_outline_generation(self):
        """测试大纲生成"""
        from core.outline_generator import OutlineGenerator

        llm = MockLLMClient()
        gen = OutlineGenerator(llm)

        outline = gen.generate(
            topic="基于 Spring Boot 的码头船只出行管理系统设计与实现",
            keywords=["Spring Boot", "码头管理", "船只调度"],
            discipline="软件工程",
        )

        self.assertIsNotNone(outline)
        self.assertIsNotNone(outline.root)
        chapters = outline.get_chapters()
        self.assertGreater(len(chapters), 3, "大纲应至少包含 4 章")
        self.assertIn("绪论", chapters[0].title)

        print(f"  ✅ 大纲生成: {len(chapters)} 章")

    def test_02_chapter_generation(self):
        """测试逐章生成"""
        from core.outline_generator import OutlineGenerator
        from core.chapter_generator import ChapterGenerator

        llm = MockLLMClient()
        og = OutlineGenerator(llm)
        outline = og.generate(
            topic="基于 Spring Boot 的码头船只出行管理系统",
            keywords=["Spring Boot"],
            discipline="软件工程",
        )

        cg = ChapterGenerator(llm)
        chapters = []

        chapter_nodes = outline.get_chapters()
        for node in chapter_nodes:
            ch = cg.generate_chapter(outline, node, chapters)
            self.assertIsNotNone(ch.content_markdown)
            self.assertGreater(estimate_word_count(ch.content_markdown), 100)
            self.assertEqual(ch.status, "done")
            chapters.append(ch)

        self.assertEqual(len(chapters), len(chapter_nodes))
        total_words = sum(ch.word_count for ch in chapters)
        print(f"  ✅ 章节生成: {len(chapters)} 章, {total_words:,} 字")

    def test_03_thesis_assembly(self):
        """测试 Thesis 对象组装"""
        from core.outline_generator import OutlineGenerator

        llm = MockLLMClient()
        og = OutlineGenerator(llm)
        outline = og.generate(
            topic="基于 Spring Boot 的码头船只出行管理系统",
            keywords=["Spring Boot"],
            discipline="软件工程",
        )

        chapter_nodes = outline.get_chapters()
        chapters = [Chapter(node=node, status="done") for node in chapter_nodes]

        for i, ch in enumerate(chapters):
            ch.content_markdown = MOCK_CHAPTER_CONTENTS[i]
            ch.word_count = estimate_word_count(ch.content_markdown)

        thesis = Thesis(
            topic="基于 Spring Boot 的码头船只出行管理系统设计与实现",
            keywords=["Spring Boot", "码头管理"],
            target_word_count=15000,
            outline=outline,
            chapters=chapters,
            references=MOCK_REFERENCES,
            template_styles=TemplateStyles(),
            template_path="",
        )

        self.assertEqual(len(thesis.chapters), len(chapter_nodes))
        self.assertGreater(thesis.get_total_words(), 1000)  # Mock 内容较短，降低阈值
        self.assertEqual(len(thesis.references), 3)
        print(f"  ✅ Thesis 组装: {thesis.get_total_words():,} 字, {len(thesis.references)} 篇参考文献")

    def test_04_docx_output(self):
        """测试 DOCX 输出"""
        from core.outline_generator import OutlineGenerator

        llm = MockLLMClient()
        og = OutlineGenerator(llm)
        outline = og.generate(
            topic="基于 Spring Boot 的码头船只出行管理系统",
            keywords=["Spring Boot"],
            discipline="软件工程",
        )

        chapter_nodes = outline.get_chapters()
        chapters = [Chapter(node=node, status="done") for node in chapter_nodes]
        for i, ch in enumerate(chapters):
            ch.content_markdown = MOCK_CHAPTER_CONTENTS[i]
            ch.word_count = estimate_word_count(ch.content_markdown)

        thesis = Thesis(
            topic="基于 Spring Boot 的码头船只出行管理系统设计与实现",
            keywords=["Spring Boot", "码头管理"],
            target_word_count=15000,
            outline=outline,
            chapters=chapters,
            references=MOCK_REFERENCES,
            template_styles=TemplateStyles(),
        )

        output_path = str(self.output_dir / "test_thesis.docx")
        formatter = DocxFormatter()
        result_path = formatter.create_document(thesis, output_path)

        self.assertTrue(os.path.exists(result_path))
        file_size = os.path.getsize(result_path)
        self.assertGreater(file_size, 1024, "DOCX 文件应大于 1KB")

        # 验证 DOCX 结构
        from docx import Document
        doc = Document(result_path)
        paragraphs = [p.text for p in doc.paragraphs]
        full_text = "\n".join(paragraphs)

        # 关键标题
        self.assertIn("基于 Spring Boot 的码头船只出行管理系统设计与实现", full_text)
        self.assertIn("绪论", full_text)

        # 表格
        tables = doc.tables
        self.assertGreater(len(tables), 0, "应至少包含一个三线表")

        print(f"  ✅ DOCX 输出: {result_path} ({file_size:,} bytes, {len(tables)} tables)")

    def test_05_watermark_and_disclaimer(self):
        """测试水印和免责声明注入"""
        from core.outline_generator import OutlineGenerator

        llm = MockLLMClient()
        og = OutlineGenerator(llm)
        outline = og.generate(
            topic="基于 Spring Boot 的码头船只出行管理系统",
            keywords=["Spring Boot"],
            discipline="软件工程",
        )

        chapter_nodes = outline.get_chapters()
        chapters = [Chapter(node=node, status="done") for node in chapter_nodes]
        for i, ch in enumerate(chapters):
            ch.content_markdown = MOCK_CHAPTER_CONTENTS[i]
            ch.word_count = estimate_word_count(ch.content_markdown)

        thesis = Thesis(
            topic="测试论文",
            keywords=["测试"],
            target_word_count=5000,
            outline=outline,
            chapters=chapters,
            references=MOCK_REFERENCES[:1],
            template_styles=TemplateStyles(),
        )

        raw_path = str(self.output_dir / "raw.docx")
        final_path = str(self.output_dir / "final.docx")

        formatter = DocxFormatter()
        formatter.create_document(thesis, raw_path)

        add_watermark_and_disclaimer(raw_path, final_path)
        self.assertTrue(os.path.exists(final_path))

        doc = Document(final_path)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("免责声明", full_text)
        print(f"  ✅ 水印与免责声明: {final_path}")

    def test_06_full_pipeline_integration(self):
        """完整链路集成测试"""
        from core.outline_generator import OutlineGenerator
        from core.chapter_generator import ChapterGenerator
        from utils.watermark import add_watermark_and_disclaimer

        # Phase 1: 生成大纲
        llm = MockLLMClient()
        og = OutlineGenerator(llm)
        outline = og.generate(
            topic="基于 Spring Boot 的码头船只出行管理系统设计与实现",
            keywords=["Spring Boot", "码头管理", "船只调度"],
            discipline="软件工程",
        )

        # Phase 2: 逐章生成
        cg = ChapterGenerator(llm)
        chapter_nodes = outline.get_chapters()
        chapters = []
        for node in chapter_nodes:
            ch = cg.generate_chapter(outline, node, chapters)
            chapters.append(ch)

        # Phase 3: 组装 Thesis
        thesis = Thesis(
            topic="基于 Spring Boot 的码头船只出行管理系统设计与实现",
            keywords=["Spring Boot", "码头管理", "船只调度"],
            target_word_count=15000,
            outline=outline,
            chapters=chapters,
            references=MOCK_REFERENCES,
            template_styles=TemplateStyles(),
        )

        # Phase 4: DOCX
        output_path = str(self.output_dir / "full_pipeline.docx")
        formatter = DocxFormatter()
        formatter.create_document(thesis, output_path)

        # Phase 5: 水印
        final_path = str(self.output_dir / "full_pipeline_final.docx")
        add_watermark_and_disclaimer(output_path, final_path)

        # 验证
        self.assertTrue(os.path.exists(final_path))
        file_size = os.path.getsize(final_path)
        self.assertGreater(file_size, 2048)

        doc = Document(final_path)
        full_text = "\n".join(p.text for p in doc.paragraphs)

        # 结构验证
        required_sections = [
            "绪论",
            "研究背景",
            "技术基础",
            "需求分析",
            "系统设计",
            "系统实现",
            "总结",
            "参考文献",
            "免责声明",
        ]
        found_sections = [s for s in required_sections if s in full_text]
        print(f"  ✅ 完整链路: {final_path} ({file_size:,} bytes)")
        print(f"     结构检查: {len(found_sections)}/{len(required_sections)} sections found")
        print(f"     缺失: {set(required_sections) - set(found_sections)}")

        self.assertGreater(len(found_sections), len(required_sections) * 0.7,
                          f"应至少包含 70% 的关键章节，实际 {len(found_sections)}/{len(required_sections)}")

    @classmethod
    def tearDownClass(cls):
        """清理临时文件"""
        import shutil
        try:
            shutil.rmtree(cls.temp_dir)
        except Exception:
            pass


if __name__ == "__main__":
    print("=" * 60)
    print("  Mock 端到端测试 — 核心生成链路验证")
    print("  测试主题: 基于 Spring Boot 的码头船只出行管理系统")
    print("  无需 API Key，使用预制 LLM 响应")
    print("=" * 60)
    print()

    # 运行测试
    loader = unittest.TestLoader()
    suite = loader.loadTestsFromTestCase(TestE2EMock)
    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(suite)

    # 总结
    print(f"\n{'=' * 60}")
    if result.wasSuccessful():
        print(f"  🎉 全部 {result.testsRun} 个端到端测试通过")
        print(f"  核心生成链路: 主题 → 大纲 → 6章 → DOCX → 水印 → ✅")
    else:
        print(f"  ❌ {len(result.failures)} 个失败, {len(result.errors)} 个错误")
        for test, traceback in result.failures + result.errors:
            print(f"\n  --- {test} ---")
            print(traceback[-500:])
    print(f"{'=' * 60}")
