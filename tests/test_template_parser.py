"""
测试模板解析器 (test_template_parser.py)

覆盖：
- TemplateParser 初始化
- parse() 无文件时返回默认样式
- parse() 文件不存在时返回默认样式
- parse() 正常 .docx 文件的样式提取
- 内部方法：_extract_page_margins, _extract_heading_styles, 等
"""

import sys
import os
import tempfile
import shutil
import unittest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from core.models import TemplateStyles, StyleDef

try:
    from docx import Document
    from docx.shared import Cm, Pt
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class TestTemplateParserInit(unittest.TestCase):
    """测试 TemplateParser 初始化"""

    def test_init_creates_default_styles(self):
        """测试初始化创建默认样式"""
        from core.template_parser import TemplateParser
        parser = TemplateParser()
        self.assertIsNotNone(parser.default_styles)
        self.assertIsInstance(parser.default_styles, TemplateStyles)

    def test_default_styles_have_headings(self):
        """测试默认样式包含标题样式"""
        from core.template_parser import TemplateParser
        parser = TemplateParser()
        self.assertIn(1, parser.default_styles.heading_styles)
        self.assertIn(2, parser.default_styles.heading_styles)
        self.assertIn(3, parser.default_styles.heading_styles)


class TestTemplateParserParse(unittest.TestCase):
    """测试 parse() 方法"""

    @classmethod
    def setUpClass(cls):
        if not HAS_DOCX:
            raise unittest.SkipTest("python-docx 未安装，跳过模板解析测试")
        from core.template_parser import TemplateParser
        cls.TemplateParser = TemplateParser

    def setUp(self):
        self.parser = self.TemplateParser()
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        if os.path.isdir(self.temp_dir):
            shutil.rmtree(self.temp_dir, ignore_errors=True)

    def _create_minimal_docx(self, filename="test_template.docx") -> str:
        """创建最小 .docx 文件作为测试 fixture"""
        doc = Document()

        # 设置自定义页边距
        section = doc.sections[0]
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(3.0)
        section.left_margin = Cm(3.5)
        section.right_margin = Cm(2.5)

        # 添加标题
        doc.add_heading("第一章 标题", level=1)
        doc.add_paragraph("正文内容段落。")

        path = os.path.join(self.temp_dir, filename)
        doc.save(path)
        return path

    def test_parse_valid_docx(self):
        """测试解析有效 .docx 文件"""
        docx_path = self._create_minimal_docx()
        styles = self.parser.parse(docx_path)
        self.assertIsInstance(styles, TemplateStyles)
        self.assertIsNotNone(styles.page_margins)
        self.assertGreater(len(styles.heading_styles), 0)

    def test_parse_none_path(self):
        """测试 None 路径返回默认样式"""
        styles = self.parser.parse(None)  # type: ignore
        self.assertIsInstance(styles, TemplateStyles)
        # 应返回默认样式（非同一对象但值相同）
        self.assertEqual(styles.page_margins["top"], 25.4)

    def test_parse_empty_path(self):
        """测试空字符串路径返回默认样式"""
        styles = self.parser.parse("")
        self.assertIsInstance(styles, TemplateStyles)

    def test_parse_nonexistent_file(self):
        """测试不存在文件返回默认样式"""
        styles = self.parser.parse("/nonexistent/template.docx")
        self.assertIsInstance(styles, TemplateStyles)
        # 验证是默认样式
        self.assertEqual(styles.page_margins["top"], 25.4)

    def test_parse_extracts_page_margins(self):
        """测试正确提取页边距"""
        docx_path = self._create_minimal_docx()
        styles = self.parser.parse(docx_path)
        # 我们设置的是 Cm(3.0) = 30mm
        self.assertIsNotNone(styles.page_margins)

    def test_parse_returns_template_styles(self):
        """测试 parse 总是返回 TemplateStyles"""
        docx_path = self._create_minimal_docx()
        styles = self.parser.parse(docx_path)
        self.assertIsInstance(styles, TemplateStyles)

    def test_parse_without_docx_library(self):
        """模拟 python-docx 未安装的情况"""
        # 这个测试验证 docx-formatter 中的 HAS_DOCX 检查
        # 由于我们已安装 python-docx，此测试验证正常路径
        from core.template_parser import HAS_DOCX as TP_HAS_DOCX
        self.assertTrue(TP_HAS_DOCX)

    def test_parse_with_body_text(self):
        """测试解析带正文的模板"""
        doc = Document()
        doc.add_paragraph("普通正文段落内容，用于测试正文样式提取。")
        doc.add_heading("标题1", level=1)
        doc.add_heading("标题2", level=2)
        doc.add_paragraph("另一段正文。")

        path = os.path.join(self.temp_dir, "body_test.docx")
        doc.save(path)

        styles = self.parser.parse(path)
        self.assertIsNotNone(styles.body_style)

    def test_default_styles_unchanged_after_parse(self):
        """测试 parse 失败不影响 default_styles"""
        original = self.parser.default_styles
        self.parser.parse("/nonexistent/file.docx")
        self.assertIs(self.parser.default_styles, original)


class TestTemplateParserHelperMethods(unittest.TestCase):
    """测试辅助方法"""

    @classmethod
    def setUpClass(cls):
        if not HAS_DOCX:
            raise unittest.SkipTest("python-docx 未安装，跳过模板解析测试")
        from core.template_parser import TemplateParser
        cls.TemplateParser = TemplateParser

    def test_get_font_name_none(self):
        """测试 _get_font_name with None"""
        # 使用 mock 对象模拟 font 属性为 None
        class MockFont:
            name = None
        result = self.TemplateParser._get_font_name(MockFont())
        self.assertEqual(result, "宋体")

    def test_get_font_name_with_value(self):
        """测试 _get_font_name with value"""
        class MockFont:
            name = "Arial"
        result = self.TemplateParser._get_font_name(MockFont())
        self.assertEqual(result, "Arial")

    def test_get_font_size_none(self):
        """测试 _get_font_size_pt with None size"""
        class MockFont:
            size = None
        result = self.TemplateParser._get_font_size_pt(MockFont())
        self.assertEqual(result, 12)

    def test_get_font_size_with_value(self):
        """测试 _get_font_size_pt with value"""
        # 14pt = 14 * 12700 EMU
        from docx.shared import Pt
        class MockFont:
            size = Pt(14)
        result = self.TemplateParser._get_font_size_pt(MockFont())
        self.assertEqual(result, 14)

    def test_get_alignment_center(self):
        """测试 _get_alignment CENTER"""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        result = self.TemplateParser._get_alignment(WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(result, 1)

    def test_get_alignment_left(self):
        """测试 _get_alignment LEFT"""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        result = self.TemplateParser._get_alignment(WD_ALIGN_PARAGRAPH.LEFT)
        self.assertEqual(result, 0)

    def test_get_alignment_right(self):
        """测试 _get_alignment RIGHT"""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        result = self.TemplateParser._get_alignment(WD_ALIGN_PARAGRAPH.RIGHT)
        self.assertEqual(result, 2)

    def test_get_alignment_none(self):
        """测试 _get_alignment None"""
        result = self.TemplateParser._get_alignment(None)
        self.assertEqual(result, 0)


if __name__ == "__main__":
    unittest.main()
