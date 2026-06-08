"""
水印与免责声明注入模块

在生成的 .docx 文件中添加：
1. 每页水印文字（"AI 辅助生成 · 仅供参考"）
2. 文档末尾的免责声明页面

伦理护栏：确保所有输出文档明确标注 AI 辅助生成属性。
"""

import os
import logging
from typing import Optional, Dict, List

logger = logging.getLogger(__name__)

try:
    from lxml import etree
    HAS_LXML = True
except ImportError:
    etree = None  # type: ignore
    HAS_LXML = False

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    HAS_DOCX = True
except ImportError:
    Document = None  # type: ignore
    HAS_DOCX = False


# 免责声明文本
DISCLAIMER_TITLE = "免责声明"
DISCLAIMER_BODY = (
    "本文档由「论文自动生成系统」使用 AI 大语言模型辅助生成，仅供学习、研究和参考之用。\n\n"
    "1. 本系统生成的论文内容基于用户输入的主题和大纲，由 AI 自动撰写，"
    "不代表任何学术机构的观点或立场。\n"
    "2. 生成内容可能存在事实性错误、逻辑不严谨或引用不当等问题，"
    "使用者应自行核实所有内容的准确性。\n"
    "3. 本系统不承担因使用生成内容而产生的任何学术不端责任。"
    "严禁将生成内容直接作为学位论文提交。\n"
    "4. 使用本系统即表示您已阅读并同意上述条款。"
)

# 水印文字
WATERMARK_TEXT = "AI 辅助生成 · 仅供参考"


def create_watermark_xml(text: str = WATERMARK_TEXT) -> str:
    """生成 Word 水印的 XML 片段

    使用 VML (Vector Markup Language) 在页眉中插入水印。

    Args:
        text: 水印文字

    Returns:
        str: 水印 XML 字符串
    """
    return f'''<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <v:shapetype id="_x0000_t136" coordsize="21600,21600" o:spt="136" adj="10800"
    path="m@7,l@8,m@5,21600l@6,21600e" filled="f"
    xmlns:v="urn:schemas-microsoft-com:vml"
    xmlns:o="urn:schemas-microsoft-com:office:office">
    <v:stroke joinstyle="miter"/>
    <v:formulas>
      <v:f eqn="sum #0 0 10800"/>
      <v:f eqn="prod #0 2 1"/>
      <v:f eqn="sum 21600 0 @1"/>
      <v:f eqn="sum 0 0 @2"/>
      <v:f eqn="sum 21600 0 @3"/>
      <v:f eqn="if @0 @3 0"/>
      <v:f eqn="if @0 21600 @1"/>
      <v:f eqn="if @0 0 @2"/>
      <v:f eqn="if @0 @4 21600"/>
      <v:f eqn="mid @5 @6"/>
      <v:f eqn="mid @8 @5"/>
      <v:f eqn="mid @7 @8"/>
      <v:f eqn="mid @6 @7"/>
      <v:f eqn="sum @6 0 @5"/>
    </v:formulas>
  </v:shapetype>
  <v:shape id="watermarkShape" type="#_x0000_t136"
    style="position:absolute;margin-left:0;margin-top:0;width:450pt;height:200pt;rotation:315;z-index:-251658240;mso-position-horizontal:center;mso-position-vertical:center"
    fillcolor="#C0C0C0" stroked="f" opacity="0.25"
    xmlns:v="urn:schemas-microsoft-com:vml">
    <v:textpath style="font-family:&quot;宋体&quot;;font-size:1pt;font-weight:bold" string="{text}"/>
  </v:shape>
</w:pict>'''


def add_watermark(doc: "Document") -> None:
    """向 Word 文档添加水印

    通过操作文档的 header XML 来实现水印效果。

    Args:
        doc: python-docx Document 对象
    """
    if Document is None:
        return

    try:
        # 获取或创建第一个节的页眉
        section = doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False

        # 在页眉的段落中添加水印 VML
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 解析并添加水印 XML
        watermark_xml = create_watermark_xml(WATERMARK_TEXT)
        run = paragraph.add_run()
        run_element = run._element  # type: ignore[attr-defined]
        pict_element = parse_xml(watermark_xml)
        run_element.append(pict_element)

    except Exception as e:
        # 水印添加失败不应阻断主流程
        import logging
        logging.warning(f"添加水印失败: {e}")


def add_disclaimer_page(doc: "Document") -> None:
    """在文档末尾添加免责声明页面

    使用分页符开始新页面，居中显示声明标题和正文。

    Args:
        doc: python-docx Document 对象
    """
    if Document is None:
        return

    try:
        # 添加分页符
        doc.add_page_break()

        # 空行留白
        for _ in range(6):
            doc.add_paragraph()

        # 标题
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(DISCLAIMER_TITLE)
        title_run.bold = True
        title_run.font.size = Pt(18)
        title_run.font.name = "黑体"
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')  # type: ignore[attr-defined]

        # 空行
        doc.add_paragraph()

        # 正文
        body_para = doc.add_paragraph()
        body_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        body_run = body_para.add_run(DISCLAIMER_BODY)
        body_run.font.size = Pt(11)
        body_run.font.name = "宋体"
        body_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # type: ignore[attr-defined]

        # 设置正文段落首行缩进
        body_para.paragraph_format.first_line_indent = Cm(0.74)

    except Exception as e:
        import logging
        logging.warning(f"添加免责声明页失败: {e}")


def add_watermark_and_disclaimer(docx_path: str, output_path: Optional[str] = None) -> str:
    """为 .docx 文件添加水印和免责声明

    这是对外暴露的主入口函数。

    Args:
        docx_path: 输入 .docx 文件路径
        output_path: 输出路径，默认覆盖原文件

    Returns:
        str: 输出文件路径
    """
    if Document is None:
        raise ImportError("python-docx 未安装，无法处理文档")

    output_path = output_path or docx_path

    doc = Document(docx_path)
    add_watermark(doc)
    add_disclaimer_page(doc)
    doc.save(output_path)

    return output_path


def inject_ethical_header(text: str) -> str:
    """在 Markdown 文本开头注入伦理声明注释

    适用于在预览阶段提醒用户。

    Args:
        text: 原始 Markdown 文本

    Returns:
        str: 注入声明后的文本
    """
    header = (
        "<!--\n"
        "  ⚠️ 伦理声明：本文档由 AI 辅助生成，仅供学习和研究参考。\n"
        "  严禁直接作为学位论文提交。使用者应自行核实所有内容的准确性。\n"
        "-->\n\n"
    )
    return header + text


# ══════════════════════════════════════════════════════════════
# P4 新增：AI 参与度标记 (段落级)
# ══════════════════════════════════════════════════════════════

class AIParticipationLevel:
    """AI 参与度层级"""
    LEVEL_A = "A"  # 完全由 LLM 生成，用户未编辑
    LEVEL_B = "B"  # LLM 生成初稿，用户在编辑器内修改过
    LEVEL_C = "C"  # 用户在编辑器内直接撰写

    @classmethod
    def label(cls, level: str) -> str:
        """获取层级描述"""
        return {
            cls.LEVEL_A: "AI 生成，未经人工修改",
            cls.LEVEL_B: "AI 生成初稿，人工已修改",
            cls.LEVEL_C: "人工撰写",
        }.get(level, "未知")


class AIParticipationTracker:
    """追踪每个章节的 AI 参与度

    在章节生成和编辑过程中记录 AI 参与层级。
    """

    def __init__(self):
        self._marks: Dict[str, str] = {}  # {chapter_id: level}

    def mark_generated(self, chapter_id: str) -> None:
        """标记为 AI 生成（层级 A）"""
        self._marks[chapter_id] = AIParticipationLevel.LEVEL_A

    def mark_edited(self, chapter_id: str) -> None:
        """标记为人工修改（层级 B）"""
        self._marks[chapter_id] = AIParticipationLevel.LEVEL_B

    def mark_manual(self, chapter_id: str) -> None:
        """标记为人工撰写（层级 C）"""
        self._marks[chapter_id] = AIParticipationLevel.LEVEL_C

    def get_level(self, chapter_id: str) -> str:
        """获取指定章节的参与度层级"""
        return self._marks.get(chapter_id, AIParticipationLevel.LEVEL_C)

    def get_summary(self) -> str:
        """生成 AI 参与度摘要报告"""
        if not self._marks:
            return "📝 暂无 AI 参与度记录"

        total = len(self._marks)
        a_count = sum(1 for v in self._marks.values() if v == AIParticipationLevel.LEVEL_A)
        b_count = sum(1 for v in self._marks.values() if v == AIParticipationLevel.LEVEL_B)
        c_count = sum(1 for v in self._marks.values() if v == AIParticipationLevel.LEVEL_C)

        ai_ratio = (a_count + b_count) / max(total, 1) * 100

        lines = [
            "## 🤖 AI 参与度报告",
            "",
            f"| 层级 | 描述 | 章节数 |",
            f"|------|------|--------|",
            f"| A | AI 生成，未修改 | {a_count} 章 |",
            f"| B | AI 初稿，已修改 | {b_count} 章 |",
            f"| C | 人工撰写 | {c_count} 章 |",
            f"| **总计** | | **{total} 章** |",
            "",
            f"**AI 参与度**：{ai_ratio:.0f}%",
        ]

        if ai_ratio > 80:
            lines.append("\n⚠️ **AI 参与度较高**，建议增加人工修改和审核。")
        elif ai_ratio > 50:
            lines.append("\n📊 **AI 参与度适中**，请确保关键章节有人工审核。")
        else:
            lines.append("\n✅ **AI 参与度较低**，论文整体以人工写作为主。")

        for ch_id, level in self._marks.items():
            lines.append(f"- 章节 `{ch_id}`: 层级 {level} ({AIParticipationLevel.label(level)})")

        return "\n".join(lines)


def inject_ai_participation_text(thesis_data: Dict, show_marks: bool = True) -> str:
    """在 DOCX 前言中注入 AI 参与度声明（P4 新增）

    Args:
        thesis_data: 论文数据字典 {topic, chapters: [{id, status, title}]}
        show_marks: 是否显示详细标记

    Returns:
        str: 格式化的声明文本
    """
    tracker = AIParticipationTracker()
    chapters = thesis_data.get("chapters", [])
    for ch in chapters:
        ch_id = ch.get("id", "")
        status = ch.get("status", "pending")
        if status == "edited":
            tracker.mark_edited(ch_id)
        elif status in ("done",):
            tracker.mark_generated(ch_id)
        else:
            tracker.mark_manual(ch_id)

    return tracker.get_summary() if show_marks else ""

